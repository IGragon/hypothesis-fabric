from __future__ import annotations

import os
from unittest.mock import MagicMock, patch

import pytest
from fastapi.testclient import TestClient


@pytest.fixture
def client(tmp_path, monkeypatch):
    import hfabric.orchestrator.wiring as wiring
    from hfabric.api import app as app_mod

    monkeypatch.chdir(tmp_path)
    app_mod._sessions.clear()

    std = tmp_path / "additional_material"
    ex = std / "Пример 1"
    ex.mkdir(parents=True)

    import docx

    doc = docx.Document()
    doc.add_paragraph("Гипотеза 1: повысить извлечение Au на 5%.")
    doc.add_paragraph("Обоснование: ксантанат улучшает флотацию.")
    doc.save(str(ex / "Гипотезы КГМК.docx"))
    (ex / "Хвосты КГМК.xlsx").write_bytes(b"placeholder")

    with patch.object(app_mod, "_build_session_index", lambda *a, **k: None), \
         patch.object(wiring, "build_real_orchestrator", return_value=MagicMock()):
        yield TestClient(app_mod.app)


class TestExamplesOutput:
    def test_examples_contain_output_text(self, client, tmp_path):
        r = client.get("/examples")
        assert r.status_code == 200
        examples = r.json()["examples"]
        assert len(examples) == 1
        ex = examples[0]
        assert ex["name"] == "Пример 1"
        assert ex["output_file"] == "Гипотезы КГМК.docx"
        assert "Гипотеза 1" in ex["output_text"]
        assert "ксантанат" in ex["output_text"]

    def test_examples_output_text_safe_when_docx_missing(self, tmp_path, monkeypatch):
        import hfabric.orchestrator.wiring as wiring
        from hfabric.api import app as app_mod

        monkeypatch.chdir(tmp_path)
        std = tmp_path / "additional_material" / "Пример 2"
        std.mkdir(parents=True)
        (std / "Хвосты.xlsx").write_bytes(b"x")

        app_mod._sessions.clear()
        with patch.object(app_mod, "_build_session_index", lambda *a, **k: None), \
             patch.object(wiring, "build_real_orchestrator", return_value=MagicMock()):
            c = TestClient(app_mod.app)
        r = c.get("/examples")
        ex = r.json()["examples"][0]
        assert ex["output_text"] == ""
        assert ex["output_file"] == ""