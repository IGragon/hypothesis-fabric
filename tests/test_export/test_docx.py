from __future__ import annotations

import os
import tempfile

import pytest

from hfabric.export.docx_writer import write_docx
from hfabric.schemas import (
    EvidenceChunk,
    ExplainedHypothesis,
    Hypothesis,
    KPI,
    KPIParsed,
    RunResult,
    ScoredHypothesis,
)


@pytest.fixture
def temp_workdir(monkeypatch):
    with tempfile.TemporaryDirectory() as tmpdir:
        monkeypatch.chdir(tmpdir)
        yield tmpdir


@pytest.fixture
def sample_run_result():
    kpi = KPIParsed(
        goal="increase Au flotation recovery by 5% without raising cyanide use",
        kpi=KPI(metric="Au recovery", direction="increase", target="5%"),
        constraints=["no cyanide increase"],
        language="en",
    )

    evidence = {
        "chunk_001": EvidenceChunk(
            chunk_id="chunk_001",
            doc_id="doc_1",
            text="Xanthate collectors improve gold flotation recovery.",
            meta={"source": "kb"},
        ),
        "chunk_002": EvidenceChunk(
            chunk_id="chunk_002",
            doc_id="doc_2",
            text="Sodium sulphide activates oxidized gold ores.",
            meta={"source": "session"},
        ),
    }

    ranked = [
        ExplainedHypothesis(
            scored=ScoredHypothesis(
                hypothesis=Hypothesis(
                    claim="Xanthate collector addition increases Au recovery",
                    mechanism="Xanthates chemisorb on gold surfaces increasing hydrophobicity",
                    expected_effect="+5-10% Au recovery",
                    evidence_refs=["chunk_001"],
                ),
                score=0.85,
                features={"novelty": 0.5, "feasibility": 0.8, "effect": 0.9},
                cited_refs={"chunk_001": evidence["chunk_001"]},
            ),
            justification="Strong supporting evidence from literature.",
            uncertainty="Limited to lab-scale results.",
            verification_plan="Run pilot flotation test with 50g/t xanthate.",
            graph_neighbourhood=["gold --has_property--> Au recovery", "xanthate --improves--> gold"],
        ),
        ExplainedHypothesis(
            scored=ScoredHypothesis(
                hypothesis=Hypothesis(
                    claim="Sodium sulphide pre-treatment activates oxidized gold",
                    mechanism="Sulphidization forms a hydrophobic layer on oxidized gold particles",
                    expected_effect="+3-7% Au recovery for oxidized ores",
                    evidence_refs=["chunk_002"],
                ),
                score=0.72,
                features={"novelty": 0.6, "feasibility": 0.7, "effect": 0.7},
                cited_refs={"chunk_002": evidence["chunk_002"]},
            ),
            justification="Relevant for oxidized ores in current circuit.",
            uncertainty="Sulphidization dosage needs optimization.",
            verification_plan="Test Na2S at 100-500g/t on oxide samples.",
            graph_neighbourhood=["sodium_sulphide --activates--> oxidized_gold"],
        ),
    ]

    return RunResult(
        run_id="test-run-001",
        session_id="test-session",
        query="How can we increase Au flotation recovery?",
        kpi=kpi,
        ranked=ranked,
        export_path=None,
        status="complete",
    )


class TestDocxWriter:
    def test_creates_file(self, temp_workdir, sample_run_result):
        path = write_docx(sample_run_result, sample_run_result.session_id)

        assert os.path.exists(path)
        assert os.path.getsize(path) > 0
        assert path.endswith("report.docx")

    def test_path_created_correctly(self, temp_workdir, sample_run_result):
        path = write_docx(sample_run_result, "session-abc")

        expected_dir = os.path.join("sessions", "session-abc", "export")
        assert os.path.isdir(expected_dir)
        assert path == os.path.join(expected_dir, "report.docx")

    def test_returns_string_path(self, temp_workdir, sample_run_result):
        path = write_docx(sample_run_result, sample_run_result.session_id)

        assert isinstance(path, str)
        assert len(path) > 0

    def test_empty_ranked_list(self, temp_workdir, sample_run_result):
        result = sample_run_result.model_copy(deep=True)
        result.ranked = []

        path = write_docx(result, result.session_id)
        assert os.path.exists(path)
        assert os.path.getsize(path) > 0

    def test_idempotent(self, temp_workdir, sample_run_result):
        path1 = write_docx(sample_run_result, sample_run_result.session_id)
        path2 = write_docx(sample_run_result, sample_run_result.session_id)

        assert path1 == path2
        assert os.path.exists(path1)
        assert os.path.exists(path2)

    def test_kpi_target_none(self, temp_workdir, sample_run_result):
        result = sample_run_result.model_copy(deep=True)
        result.kpi.kpi.target = None

        path = write_docx(result, result.session_id)
        assert os.path.exists(path)
        assert os.path.getsize(path) > 0

    def test_no_constraints(self, temp_workdir, sample_run_result):
        result = sample_run_result.model_copy(deep=True)
        result.kpi.constraints = []

        path = write_docx(result, result.session_id)
        assert os.path.exists(path)
        assert os.path.getsize(path) > 0

    def test_no_evidence_refs(self, temp_workdir, sample_run_result):
        result = sample_run_result.model_copy(deep=True)
        result.ranked[0].scored.cited_refs = {}
        result.ranked[0].scored.hypothesis.evidence_refs = []

        path = write_docx(result, result.session_id)
        assert os.path.exists(path)
        assert os.path.getsize(path) > 0

    def test_no_features(self, temp_workdir, sample_run_result):
        result = sample_run_result.model_copy(deep=True)
        result.ranked[0].scored.features = {}

        path = write_docx(result, result.session_id)
        assert os.path.exists(path)
        assert os.path.getsize(path) > 0

    def test_docx_content_contains_key_fields(self, temp_workdir, sample_run_result):
        path = write_docx(sample_run_result, sample_run_result.session_id)

        from docx import Document
        doc = Document(path)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "Hypothesis Fabric Report" in full_text
        assert sample_run_result.query in full_text
        assert "Au recovery" in full_text
        assert "Ranked Hypotheses" in full_text
        assert "Xanthate collector addition increases Au recovery" in full_text
        assert "Xanthates chemisorb" in full_text
        assert "Strong supporting evidence from literature" in full_text
        assert "Limited to lab-scale results" in full_text
        assert "Run pilot flotation test" in full_text

    def test_docx_contains_two_hypotheses(self, temp_workdir, sample_run_result):
        path = write_docx(sample_run_result, sample_run_result.session_id)

        from docx import Document
        doc = Document(path)
        heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        hypothesis_headings = [h for h in heading_texts if h.startswith("Hypothesis ")]
        assert len(hypothesis_headings) == 2
