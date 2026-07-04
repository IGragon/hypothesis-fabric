from __future__ import annotations

import os

from hfabric.schemas import RunResult


def write_docx(result: RunResult, session_id: str) -> str:
    from docx import Document

    export_dir = os.path.join("sessions", session_id, "export")
    os.makedirs(export_dir, exist_ok=True)
    path = os.path.join(export_dir, "report.docx")

    doc = Document()

    doc.add_heading("Hypothesis Fabric Report", level=0)

    doc.add_heading("Query", level=1)
    doc.add_paragraph(result.query)

    doc.add_heading("KPI Summary", level=1)
    kpi = result.kpi.kpi
    doc.add_paragraph(f"Metric: {kpi.metric}")
    doc.add_paragraph(f"Direction: {kpi.direction}")
    target_str = kpi.target if kpi.target else "N/A"
    doc.add_paragraph(f"Target: {target_str}")

    constraints = result.kpi.constraints
    if constraints:
        doc.add_heading("Constraints", level=2)
        for c in constraints:
            doc.add_paragraph(c, style="List Bullet")

    doc.add_heading("Ranked Hypotheses", level=1)

    for i, eh in enumerate(result.ranked or [], 1):
        h = eh.scored.hypothesis
        doc.add_heading(f"Hypothesis {i}", level=2)

        doc.add_heading("Claim", level=3)
        doc.add_paragraph(h.claim)

        doc.add_heading("Mechanism", level=3)
        doc.add_paragraph(h.mechanism)

        doc.add_heading("Expected Effect", level=3)
        doc.add_paragraph(h.expected_effect)

        doc.add_heading("Score", level=3)
        doc.add_paragraph(f"{eh.scored.score:.3f}")

        features = eh.scored.features
        if features:
            doc.add_heading("Features", level=3)
            for fname, fval in features.items():
                doc.add_paragraph(f"{fname}: {fval:.3f}")

        doc.add_heading("Justification", level=3)
        doc.add_paragraph(eh.justification)

        doc.add_heading("Novelty", level=3)
        doc.add_paragraph(eh.novelty)

        doc.add_heading("Risks", level=3)
        doc.add_paragraph(eh.risks)

        doc.add_heading("Why it matters (value / KPI)", level=3)
        doc.add_paragraph(eh.why_it_matters)

        if eh.effect_cause_examples:
            doc.add_heading("Effect-cause examples", level=3)
            for ex in eh.effect_cause_examples:
                doc.add_paragraph(ex, style="List Bullet")

        doc.add_heading("General approach", level=3)
        doc.add_paragraph(eh.general_approach)

        doc.add_heading("What to do now", level=3)
        doc.add_paragraph(eh.actionable_now)

        doc.add_heading("Existing best practices", level=3)
        doc.add_paragraph(eh.best_practices)

        doc.add_heading("Uncertainty", level=3)
        doc.add_paragraph(eh.uncertainty)

        doc.add_heading("Verification Plan (roadmap)", level=3)
        doc.add_paragraph(eh.verification_plan)

        cited_refs = eh.scored.cited_refs
        if cited_refs:
            doc.add_heading("Evidence / Sources", level=3)
            for ref_id, chunk in cited_refs.items():
                url = chunk.meta.get("url")
                if url:
                    doc.add_paragraph(f"[{ref_id}] {chunk.meta.get('title') or url} — {url}")
                else:
                    doc.add_paragraph(f"[{ref_id}] {chunk.text[:300]}")

        doc.add_paragraph()

    doc.save(path)
    return path
