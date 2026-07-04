from __future__ import annotations

import base64
import os
from typing import Any

IMAGE_EXTS = (".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".webp")
_PDF_MIN_TEXT_CHARS = 50
_PDF_IMAGE_PAGES_MAX = 4

_LOCAL_TESSERACT = "/tmp/opencode/tesseract_local/usr/bin/tesseract"
_LOCAL_TESSDATA = "/tmp/opencode/tesseract_local/usr/share/tesseract-ocr/5/tessdata"
_LOCAL_LIB = "/tmp/opencode/tesseract_local/usr/lib/x86_64-linux-gnu"


def _setup_tesseract() -> None:
    if os.path.isfile(_LOCAL_TESSERACT):
        os.environ.setdefault("TESSDATA_PREFIX", _LOCAL_TESSDATA)
        os.environ.setdefault("LD_LIBRARY_PATH", _LOCAL_LIB)
        try:
            import pytesseract

            if not pytesseract.pytesseract.tesseract_cmd or pytesseract.pytesseract.tesseract_cmd == "tesseract":
                pytesseract.pytesseract.tesseract_cmd = _LOCAL_TESSERACT
        except Exception:
            pass


_setup_tesseract()


def _doc_id(path: str) -> str:
    return os.path.splitext(os.path.basename(path))[0]


def _run_ocr(path_or_image: Any, enabled: bool = True) -> str:
    if not enabled:
        return ""
    try:
        import pytesseract
        from PIL import Image

        img = path_or_image
        if isinstance(path_or_image, str):
            img = Image.open(path_or_image)
        try:
            return pytesseract.image_to_string(img, lang="rus+eng").strip()
        except Exception:
            return pytesseract.image_to_string(img).strip()
    except Exception:
        pass
    return _run_ocr_pymupdf(path_or_image)


def _run_ocr_pymupdf(path_or_image: Any) -> str:
    try:
        import fitz

        doc = None
        if isinstance(path_or_image, str):
            doc = fitz.open(path_or_image)
            page = doc[0]
        else:
            import io

            from PIL import Image as PILImage

            buf = io.BytesIO()
            img = path_or_image
            if isinstance(path_or_image, str):
                img = PILImage.open(path_or_image)
            img.save(buf, format="PNG")
            doc = fitz.open(stream=buf.getvalue(), filetype="png")
            page = doc[0]
        tp = page.get_textpage_ocr(flags=0, language="rus+eng", dpi=150, full=True)
        text = tp.get_text().strip()
        if doc is not None:
            doc.close()
        return text
    except Exception:
        return ""


def _run_vlm(image_bytes: bytes, config: Any, vision_model: Any, mime: str = "image/png") -> str:
    if config is not None and not getattr(config, "enable_vlm", True):
        return ""
    model = vision_model
    if model is None:
        if config is None:
            return ""
        from hfabric.llm import create_vision_chat_model

        model = create_vision_chat_model(config)
    if model is None:
        return ""
    try:
        from hfabric.llm import vlm_describe_image

        b64 = base64.b64encode(image_bytes).decode("ascii")
        return vlm_describe_image(model, b64, mime=mime)
    except Exception:
        return ""


def parse_image(path: str, config: Any = None, vision_model: Any = None) -> list[dict]:
    doc_id = _doc_id(path)
    ocr_enabled = config is None or getattr(config, "enable_ocr", True)
    ocr_text = _run_ocr(path, enabled=ocr_enabled)

    lower = path.lower()
    is_schematic = "схем" in lower or "регламент" in lower or "оборудован" in lower
    need_vlm = len(ocr_text) < 40 or is_schematic

    vlm_description = ""
    if need_vlm:
        try:
            with open(path, "rb") as fh:
                image_bytes = fh.read()
            ext = os.path.splitext(path)[1].lower().lstrip(".")
            mime = f"image/{'jpeg' if ext in ('jpg', 'jpeg') else (ext or 'png')}"
            vlm_description = _run_vlm(image_bytes, config, vision_model, mime=mime)
        except Exception:
            vlm_description = ""

    text_parts: list[str] = []
    if vlm_description:
        text_parts.append(vlm_description)
    if ocr_text:
        text_parts.append(f"OCR: {ocr_text}")
    text = "\n".join(text_parts).strip()

    low_confidence = not bool(text)
    if not text:
        text = f"[image {os.path.basename(path)} — no extractable text]"

    return [{
        "text": text,
        "meta": {
            "page": 1,
            "path": path,
            "doc_id": doc_id,
            "image": True,
            "ocr_text": ocr_text,
            "vlm_description": vlm_description,
            "low_confidence": low_confidence,
        },
    }]


def parse_docx(path: str) -> list[dict]:
    import docx

    document = docx.Document(path)
    doc_id = _doc_id(path)
    pages: list[dict] = []
    current_section = "Документ"
    buf: list[str] = []
    counter = [1]

    def flush() -> None:
        if buf:
            pages.append({
                "text": "\n".join(buf),
                "meta": {"page": counter[0], "path": path, "doc_id": doc_id, "section": current_section},
            })
            counter[0] += 1
            buf.clear()

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if "heading" in style or "заголов" in style or "title" in style:
            flush()
            current_section = text
        buf.append(text)
    flush()

    for ti, table in enumerate(document.tables):
        rows: list[str] = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            current_section = f"Таблица {ti + 1}"
            buf.extend(rows)
            flush()

    if not pages:
        pages.append({
            "text": f"[docx {os.path.basename(path)} — no extractable text]",
            "meta": {"page": 1, "path": path, "doc_id": doc_id, "section": "Документ", "low_confidence": True},
        })
    return pages


def _format_number(v, decimals: int = 2) -> str:
    if v is None:
        return "—"
    if isinstance(v, float):
        return f"{v:.{decimals}f}"
    return str(v)


def _is_data_row(vals: list) -> bool:
    if vals[1] is None or not isinstance(vals[1], str):
        return False
    label = vals[1].strip().lower()
    skip_labels = {
        "материал", "поступило в переработку", "отвальные хвосты",
        "хвосты породные", "класс крупности, мкм",
    }
    return label not in skip_labels


def parse_pdf(path: str, config: Any = None, vision_model: Any = None) -> list[dict]:
    import fitz

    doc = fitz.open(path)
    doc_id = os.path.splitext(os.path.basename(path))[0]
    results: list[dict] = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text()
        if text.strip():
            results.append({
                "text": text,
                "meta": {"page": page_num + 1, "path": path, "doc_id": doc_id},
            })
    doc.close()
    return results


def _parse_tailings_xlsx(path: str) -> list[dict]:
    from openpyxl import load_workbook

    wb = load_workbook(path, data_only=True)
    ws = wb["Итог"]
    doc_id = os.path.splitext(os.path.basename(path))[0]
    rows_raw = [
        [ws.cell(r, c).value for c in range(1, ws.max_column + 1)]
        for r in range(1, ws.max_row + 1)
    ]

    pages: list[dict] = []
    current_section: str = "Общие сведения"
    current_text: list[str] = []
    page_num = 1
    size_class_names = {"+125", "+71", "-71 + 45", "-45 + 20", " -20 + 10", "-10", "-71+45", "-45+20", "-20+10"}

    for row in rows_raw:
        if not row or all(v is None for v in row):
            continue

        col_b = (row[1] or "").strip() if len(row) > 1 else ""
        col_b_lower = col_b.lower()

        if col_b_lower in ("поступило в переработку", "отвальные хвосты", "хвосты породные"):
            if current_text:
                pages.append({
                    "text": "\n".join(current_text),
                    "meta": {"page": page_num, "path": path, "doc_id": doc_id, "section": current_section},
                })
                page_num += 1
                current_text = []
            current_section = col_b

        if col_b_lower in ("класс крупности, мкм", "класс крупности"):
            if current_text:
                pages.append({
                    "text": "\n".join(current_text),
                    "meta": {"page": page_num, "path": path, "doc_id": doc_id, "section": current_section},
                })
                page_num += 1
                current_text = []
            current_section = "Распределение по классам крупности"

        if col_b.replace(" ", "") in size_class_names or col_b in size_class_names:
            if current_text:
                pages.append({
                    "text": "\n".join(current_text),
                    "meta": {"page": page_num, "path": path, "doc_id": doc_id, "section": current_section},
                })
                page_num += 1
                current_text = []
            current_section = f"Класс крупности {col_b}"

        if col_b_lower in ("извлекаемый металл", "итого извлекаемый металл", "итого не извлекаемый металл", "итого (проверка)"):
            pass

        if _is_data_row(row):
            parts = [f"{row[1]}: "]
            col_names = {
                2: "СМТ", 3: "Элемент 28 %", 4: "Элемент 28 т",
                5: "Элемент 29 %", 6: "Элемент 29 т",
            }
            for ci, cn in col_names.items():
                if ci < len(row) and row[ci] is not None:
                    parts.append(f"{cn}={_format_number(row[ci])}")
            current_text.append("  ".join(parts))

        mineral_forms = {
            "раскрытый pnt/cp": "Раскрытый пентландит/халькопирит (извлекаемая форма)",
            "закрытый pnt/cp": "Закрытый пентландит/халькопирит (извлекаемая форма)",
            "примесь в пирротине": "Примесь в пирротине (неизвлекаемая форма)",
            "силикатная форма/валлериит": "Силикатная форма/Валлериит (неизвлекаемая форма)",
            "пирит/другие элемент 29 сульфиды": "Пирит/другие сульфиды (неизвлекаемая форма)",
            "миллерит": "Миллерит (извлекаемая форма для Элемента 28)",
            "потери (расписать)": "Потери (прочие)",
            "свободный слот": "Свободный слот",
        }

        if col_b_lower in mineral_forms:
            line = mineral_forms[col_b_lower]
            if len(row) > 3 and row[3] is not None:
                line += f" — доля потерь Эл.28: {_format_number(row[3])}%"
            if len(row) > 4 and row[4] is not None:
                line += f", Эл.28: {_format_number(row[4])} т"
            if len(row) > 5 and row[5] is not None:
                line += f" — доля потерь Эл.29: {_format_number(row[5])}%"
            if len(row) > 6 and row[6] is not None:
                line += f", Эл.29: {_format_number(row[6])} т"
            current_text.append(line)

        if col_b_lower in ("итого (проверка)",):
            if len(row) > 3 and row[3] is not None:
                current_text.append(f"Итого проверка: Эл.28 доля={_format_number(row[3])}%, Эл.28 т={_format_number(row[4])}, Эл.29 доля={_format_number(row[5])}%, Эл.29 т={_format_number(row[6])}")

        if col_b_lower.startswith("извлекаемый металл") and "итого" not in col_b_lower:
            line = "Извлекаемый металл (потенциально можно извлечь): "
            if len(row) > 3 and row[3] is not None:
                line += f"доля Эл.28={_format_number(row[3])}%, "
            if len(row) > 4 and row[4] is not None:
                line += f"Эл.28={_format_number(row[4])} т, "
            if len(row) > 5 and row[5] is not None:
                line += f"доля Эл.29={_format_number(row[5])}%, "
            if len(row) > 6 and row[6] is not None:
                line += f"Эл.29={_format_number(row[6])} т"
            current_text.append(line)

        if col_b_lower.startswith("не извлекаемый металл") and "итого" not in col_b_lower:
            line = "Не извлекаемый металл (невозможно извлечь текущей технологией): "
            if len(row) > 3 and row[3] is not None:
                line += f"доля Эл.28={_format_number(row[3])}%, "
            if len(row) > 4 and row[4] is not None:
                line += f"Эл.28={_format_number(row[4])} т"
            current_text.append(line)

        if col_b_lower == "итого извлекаемый металл":
            current_text.append(f"ВСЕГО извлекаемый металл: доля Эл.28={_format_number(row[3])}%, Эл.28={_format_number(row[4])} т, доля Эл.29={_format_number(row[5])}%, Эл.29={_format_number(row[6])} т")
            current_text.append(f"Вывод: потенциально извлекаемая доля Элемента 28 составляет {_format_number(row[3])}% "
                              f"({_format_number(row[4])} т), что указывает на значительный резерв "
                              f"для оптимизации технологии обогащения с целью снижения потерь металла с хвостами.")

        if col_b_lower == "итого не извлекаемый металл":
            current_text.append(f"ВСЕГО не извлекаемый металл: доля Эл.28={_format_number(row[3])}%, Эл.28={_format_number(row[4])} т")
            current_text.append(f"Вывод: {_format_number(row[3])}% потерь Элемента 28 "
                              f"({_format_number(row[4])} т) связано с минеральными формами, "
                              f"не извлекаемыми текущей технологией обогащения.")

    if current_text:
        pages.append({
            "text": "\n".join(current_text),
            "meta": {"page": page_num, "path": path, "doc_id": doc_id, "section": current_section},
        })

    wb.close()
    return pages


def _detect_xlsx_type(path: str) -> str | None:
    from openpyxl import load_workbook

    wb = load_workbook(path, data_only=True)
    sheet_names = [s.lower() for s in wb.sheetnames]
    wb.close()
    if "итог" in sheet_names:
        return "tailings"
    return "generic"


def parse_xlsx(path: str) -> list[dict]:
    xlsx_type = _detect_xlsx_type(path)
    if xlsx_type == "tailings":
        return _parse_tailings_xlsx(path)
    return _parse_generic_xlsx(path)


def _parse_generic_xlsx(path: str) -> list[dict]:
    from openpyxl import load_workbook

    wb = load_workbook(path, data_only=True)
    doc_id = os.path.splitext(os.path.basename(path))[0]
    pages: list[dict] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        text_lines: list[str] = []
        for row in ws.iter_rows(values_only=True):
            non_none = [str(v) for v in row if v is not None]
            if non_none:
                text_lines.append("\t".join(non_none))
        if text_lines:
            pages.append({
                "text": "\n".join(text_lines),
                "meta": {
                    "page": wb.sheetnames.index(sheet_name) + 1,
                    "path": path,
                    "doc_id": doc_id,
                    "sheet": sheet_name,
                },
            })

    wb.close()
    return pages


_STRUCTURE_PROMPT = (
    "You are a metallurgical data analyst. Below is raw text extracted via OCR "
    "from images (schematics, flow-sheets, equipment diagrams, process regulations). "
    "The text is fragmented, contains OCR artifacts, and mixes Russian and English.\n\n"
    "Your task: restructure this into a clean, coherent technical report. "
    "Organize content into logical sections. Preserve all numbers, equipment names, "
    "reagent names, and Russian terms verbatim. Remove OCR garbage. "
    "Do NOT add information that is not in the source text.\n\n"
    "Respond in plain text with section headers prefixed by '## '."
)


def structure_ocr_pages(
    merged_text: str,
    config: Any = None,
    llm: Any = None,
    timeout_seconds: float = 60.0,
) -> str:
    if not merged_text.strip():
        return ""
    if config is not None and not getattr(config, "enable_ocr_structuring", True):
        return ""
    model = llm
    if model is None and config is not None:
        try:
            from hfabric.llm import create_chat_model

            provider = config.provider
            if hasattr(provider, "value"):
                provider = provider.value
            model = create_chat_model(provider, config.model, temperature=0.0)
        except Exception:
            return ""
    if model is None:
        return ""

    prompt = f"{_STRUCTURE_PROMPT}\n\n--- RAW OCR TEXT ---\n{merged_text}"

    import threading

    result: list[str | None] = [None]

    def _call():
        try:
            response = model.invoke(prompt)
            text = response.content if hasattr(response, "content") else str(response)
            if isinstance(text, list):
                text = "".join(
                    item.get("text", "") if isinstance(item, dict) else str(item)
                    for item in text
                )
            result[0] = text.strip()
        except Exception:
            result[0] = None

    thread = threading.Thread(target=_call, daemon=True)
    thread.start()
    thread.join(timeout=timeout_seconds)
    return result[0] or ""


def parse(path: str, config: Any = None, vision_model: Any = None) -> list[dict]:
    lower = path.lower()
    if lower.endswith(".pdf"):
        return parse_pdf(path, config=config, vision_model=vision_model)
    if lower.endswith(".xlsx"):
        return parse_xlsx(path)
    if lower.endswith(".docx"):
        return parse_docx(path)
    if lower.endswith(IMAGE_EXTS):
        return parse_image(path, config=config, vision_model=vision_model)
    return []


SUPPORTED_EXTS = (".pdf", ".xlsx", ".docx") + IMAGE_EXTS
